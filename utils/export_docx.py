from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_exam_docx(header, questions, mode="student"):
    doc = Document()

    # Header school
    p_school = doc.add_paragraph()
    run = p_school.add_run(header.get("school", "").upper())
    run.bold = True
    run.font.size = Pt(12)
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"ĐỀ KIỂM TRA {header.get('semester', '').upper()} — {header.get('subject', '').upper()}")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subheader
    sub = doc.add_paragraph()
    sub.add_run(f"{header.get('grade', '')} • Thời gian: {header.get('time', '')}").font.size = Pt(12)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note line
    if header.get("note"):
        note = doc.add_paragraph()
        note.add_run(header["note"]).font.size = Pt(11)

    doc.add_paragraph(" ")

    # Split parts: Trắc nghiệm / Tự luận
    part_mcq = [q for q in questions if q["type"] in ["MCQ", "TrueFalse", "FillBlank", "Matching"]]
    part_essay = [q for q in questions if q["type"] == "Essay"]

    # Part A: Trắc nghiệm
    if part_mcq:
        pa = doc.add_paragraph()
        ra = pa.add_run("Phần A: Trắc nghiệm")
        ra.bold = True
        doc.add_paragraph(" ")

        for i, q in enumerate(part_mcq, start=1):
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {i} ({q.get('points', 0)} điểm) — {q['type']}: ")
            run.bold = True
            p.add_run(q["prompt"])

            if q["type"] == "MCQ" and q.get("options"):
                for idx, opt in enumerate(q["options"]):
                    if opt is not None:
                        op = doc.add_paragraph()
                        op.add_run(f"{chr(65+idx)}. {opt}")
            elif q["type"] == "TrueFalse":
                doc.add_paragraph("Khoanh tròn Đúng hoặc Sai.")
            elif q["type"] == "FillBlank":
                doc.add_paragraph("Điền vào chỗ trống.")
            elif q["type"] == "Matching":
                doc.add_paragraph("Ghép cột A với cột B (giáo viên bổ sung bảng).")

    # Part B: Tự luận
    if part_essay:
        doc.add_paragraph(" ")
        pb = doc.add_paragraph()
        rb = pb.add_run("Phần B: Tự luận")
        rb.bold = True
        doc.add_paragraph(" ")

        start_index = len(part_mcq) + 1
        for j, q in enumerate(part_essay, start=start_index):
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {j} ({q.get('points', 0)} điểm) — {q['type']}: ")
            run.bold = True
            p.add_run(q["prompt"])
            doc.add_paragraph("Trình bày lời giải rõ ràng, đủ bước.")

    # Answers for teacher mode
    if mode == "teacher":
        doc.add_paragraph(" ")
        ans = doc.add_paragraph()
        r = ans.add_run("Đáp án và gợi ý lời giải")
        r.bold = True
        doc.add_paragraph(" ")

        for i, q in enumerate(questions, start=1):
            p = doc.add_paragraph()
            p.add_run(f"Câu {i}: ").bold = True
            p.add_run(f"Đáp án: {q.get('answer','')}")
            if q.get("unit"):
                p.add_run(f" ({q['unit']})")
            if q.get("explanation"):
                e = doc.add_paragraph()
                e.add_run(f"Gợi ý: {q['explanation']}")

    # Footer
    doc.add_paragraph(" ")
    ftr = doc.add_paragraph()
    ftr.add_run("— Hết —").italic = True
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
